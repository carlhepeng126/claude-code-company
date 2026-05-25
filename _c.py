import re, os, sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup
import markdown

SRC = r'C:\Users\Admin\Documents\LLM Wiki\ACT笔记-深度思考\6-storage\6-1-archive-归档\项目归档\广东申菱环境\广东申菱环境-国际化人才英语能力提升方案-调整版.md'
desktop = os.path.join(os.environ['USERPROFILE'], 'Desktop')
OUT = os.path.join(desktop, '_temp_shenling.docx')
FINAL = os.path.join(desktop, '广东申菱环境-国际化人才英语能力提升方案-调整版.docx')

FONT_BODY = 'FangSong'; FONT_HEADING = 'SimHei'; FONT_TABLE = 'SimSun'
SIZE_BODY = Pt(14); SIZE_H1 = Pt(22); SIZE_H2 = Pt(16); SIZE_H3 = Pt(15); SIZE_TABLE = Pt(10.5)
LINE_SP = Pt(26); INDENT = Pt(28)

with open(SRC, 'r', encoding='utf-8') as f:
    md_text = f.read()

html_text = markdown.markdown(md_text, extensions=['tables', 'fenced_code'])
soup = BeautifulSoup(html_text, 'html.parser')
doc = Document()

style = doc.styles['Normal']
style.font.name = FONT_BODY; style.font.size = SIZE_BODY
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
style.paragraph_format.line_spacing = LINE_SP
style.paragraph_format.space_before = Pt(0); style.paragraph_format.space_after = Pt(0)
style.paragraph_format.first_line_indent = INDENT

for s in doc.sections:
    s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.17); s.right_margin = Cm(3.17)

for lk, sz in [('Heading 1', SIZE_H1), ('Heading 2', SIZE_H2), ('Heading 3', SIZE_H3)]:
    hs = doc.styles[lk]
    hs.font.name = FONT_HEADING; hs.font.size = sz; hs.font.bold = True; hs.font.color.rgb = None
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
    hs.paragraph_format.line_spacing = Pt(sz.pt * 1.5)
    hs.paragraph_format.space_before = Pt(12); hs.paragraph_format.space_after = Pt(6)

def ar(p, t, bold=False, font_name=None, font_size=None, italic=False):
    if not t: return
    r = p.add_run(t); r.bold = bold; r.italic = italic
    r.font.name = font_name or FONT_BODY; r.font.size = font_size or SIZE_BODY
    rp = r._element.get_or_add_rPr()
    rf = OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'), font_name or FONT_BODY); rp.insert(0, rf)

def np(ind=True, al=None):
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = LINE_SP
    p.paragraph_format.first_line_indent = INDENT if ind else Pt(0)
    if al is not None: p.alignment = al
    return p

def ft(e):
    if isinstance(e, str): return e
    if e.name == 'br': return '\n'
    return ''.join(ft(c) for c in e.children)

def sb(t):
    r = []; pat = re.compile(r'\*\*(.+?)\*\*'); last = 0
    for m in pat.finditer(t):
        if m.start() > last: r.append((t[last:m.start()], False))
        r.append((m.group(1), True)); last = m.end()
    if last < len(t): r.append((t[last:], False))
    return r

def pi(pe, p, fn=None, fs=None):
    fn = fn or FONT_BODY; fs = fs or SIZE_BODY
    for c in pe.children:
        if isinstance(c, str):
            for t, b in sb(c): ar(p, t, bold=b, font_name=fn, font_size=fs)
        elif c.name in ('strong', 'b'):
            for t, b in sb(c.get_text()): ar(p, t, bold=True, font_name=fn, font_size=fs)
        elif c.name in ('em', 'i'): ar(p, c.get_text(), italic=True, font_name=fn, font_size=fs)
        elif c.name == 'code': ar(p, c.get_text(), font_name='Courier New', font_size=Pt(fs.pt * 0.85))
        elif c.name == 'br': p.add_run('\n')
        elif c.name == 'a': ar(p, c.get_text(), font_name=fn, font_size=fs)
        else: pi(c, p, fn, fs)

def pe(e):
    if e.name is None: return
    t = e.name
    if t in ('h1', 'h2', 'h3'):
        tx = e.get_text().strip()
        if tx:
            p = doc.add_heading(tx, level=int(t[1]))
            for rn in p.runs:
                rn.font.name = FONT_HEADING; rn.font.size = {1: SIZE_H1, 2: SIZE_H2, 3: SIZE_H3}[int(t[1])]
                rp = rn._element.get_or_add_rPr()
                rf = OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'), FONT_HEADING); rp.insert(0, rf)
            p.paragraph_format.first_line_indent = Pt(0)
    elif t == 'p':
        tx = e.get_text().strip()
        if not tx and not e.find(['strong', 'b', 'em', 'code']): return
        if e.get('align', '') == 'center':
            p = np(ind=False, al=WD_ALIGN_PARAGRAPH.CENTER); pi(e, p)
        elif tx.startswith('>') or (e.parent and e.parent.name == 'blockquote'):
            p = np(ind=False); p.paragraph_format.left_indent = Cm(1.0)
            for tx2, b in sb(tx.lstrip('>').lstrip(' ')): ar(p, tx2, bold=b)
        else: p = np(ind=True); pi(e, p)
    elif t == 'blockquote':
        for c in e.children:
            if c.name == 'p':
                tx = c.get_text().strip(); p = np(ind=False); p.paragraph_format.left_indent = Cm(1.0)
                for tx2, b in sb(tx): ar(p, tx2, bold=b)
            elif c.name: pe(c)
    elif t == 'table': pt(e)
    elif t == 'hr':
        p = np(ind=False); p.paragraph_format.line_spacing = Pt(6)
        pp = p._element.get_or_add_pPr(); pb = OxmlElement('w:pBdr')
        bo = OxmlElement('w:bottom')
        for a, v in [('val', 'single'), ('sz', '6'), ('space', '1'), ('color', '999999')]:
            bo.set(qn('w:' + a), v)
        pb.append(bo); pp.append(pb)
    elif t == 'ul':
        for li in e.find_all('li', recursive=False):
            tx = li.get_text().strip(); p = np(ind=False)
            p.paragraph_format.left_indent = Cm(1.5)
            for tx2, b in sb(tx): ar(p, tx2, bold=b)
    elif t == 'ol':
        for li in e.find_all('li', recursive=False):
            tx = li.get_text().strip(); p = np(ind=False)
            p.paragraph_format.left_indent = Cm(1.5)
            for tx2, b in sb(tx): ar(p, tx2, bold=b)
    elif t == 'pre':
        for ln in e.get_text().strip().split('\n'):
            p = np(ind=False); p.paragraph_format.line_spacing = Pt(14)
            ar(p, ln, font_name='Courier New', font_size=Pt(9))
    elif t in ('div', 'body'):
        for c in e.children:
            if c.name: pe(c)

def pt(te):
    rows = te.find_all('tr')
    if not rows: return
    mc = max(len(r.find_all(['td', 'th'])) for r in rows)
    if mc == 0: return
    tbl = doc.add_table(rows=len(rows), cols=mc, style='Table Grid')
    for i, r in enumerate(rows):
        cs = r.find_all(['td', 'th']); ih = bool(r.find('th')) or i == 0
        for j in range(mc):
            cl = tbl.cell(i, j); cl.paragraphs[0].clear()
            if j >= len(cs): continue
            ct = cs[j]; ps = ct.find_all('p')
            if ps:
                for pii, pr in enumerate(ps):
                    p = cl.paragraphs[0] if pii == 0 else cl.add_paragraph()
                    p.paragraph_format.line_spacing = Pt(14); p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2); p.paragraph_format.first_line_indent = Pt(0)
                    for tx, b in sb(pr.get_text().strip()):
                        rn = p.add_run(tx); rn.font.name = FONT_TABLE; rn.font.size = SIZE_TABLE; rn.bold = b or ih
                        rp = rn._element.get_or_add_rPr()
                        rf = OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'), FONT_TABLE); rp.insert(0, rf)
            else:
                tx = ft(ct).strip()
                for pii, ptn in enumerate(tx.split('\n')):
                    ptn = ptn.strip()
                    if not ptn: continue
                    pp = cl.paragraphs[0] if pii == 0 else cl.add_paragraph()
                    pp.paragraph_format.line_spacing = Pt(14); pp.paragraph_format.space_before = Pt(2)
                    pp.paragraph_format.space_after = Pt(2); pp.paragraph_format.first_line_indent = Pt(0)
                    for tx2, b in sb(ptn):
                        rn = pp.add_run(tx2); rn.font.name = FONT_TABLE; rn.font.size = SIZE_TABLE; rn.bold = b or ih
                        rp = rn._element.get_or_add_rPr()
                        rf = OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'), FONT_TABLE); rp.insert(0, rf)
    np(ind=False).paragraph_format.line_spacing = Pt(6)

bd = soup.find('body')
tg = bd if bd else soup
for e in tg.children: pe(e)
doc.save(OUT)

# Remove old file and rename
if os.path.exists(FINAL):
    os.remove(FINAL)
os.rename(OUT, FINAL)
print('Done')
